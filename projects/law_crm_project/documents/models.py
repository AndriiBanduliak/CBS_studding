"""
Модели для документооборота.
Включает хранение документов, шаблоны и генерацию документов.
"""
from django.db import models
from django.contrib.postgres.search import SearchVectorField, SearchVector
from django.contrib.postgres.indexes import GinIndex
from auditlog.registry import auditlog
from core.models import User
from case_management.models import Case


class DocumentCategory(models.Model):
    """Категории документов"""
    
    name = models.CharField(
        max_length=100,
        unique=True,
        verbose_name='Назва'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Опис'
    )
    
    parent = models.ForeignKey(
        'self',
        on_delete=models.CASCADE,
        null=True,
        blank=True,
        related_name='subcategories',
        verbose_name='Батьківська категорія'
    )
    
    class Meta:
        verbose_name = 'Категорія документів'
        verbose_name_plural = 'Категорії документів'
        ordering = ['name']
    
    def __str__(self):
        if self.parent:
            return f"{self.parent.name} / {self.name}"
        return self.name


class Document(models.Model):
    """
    Модель документа с поддержкой PostgreSQL Full-Text Search.
    """
    
    DOCUMENT_TYPE_CHOICES = [
        ('contract', 'Договір'),
        ('motion', 'Клопотання'),
        ('complaint', 'Позовна заява'),
        ('appeal', 'Апеляційна скарга'),
        ('response', 'Відповідь на позов'),
        ('court_decision', 'Рішення суду'),
        ('power_of_attorney', 'Довіреність'),
        ('evidence', 'Доказ'),
        ('correspondence', 'Листування'),
        ('other', 'Інше'),
    ]
    
    case = models.ForeignKey(
        Case,
        on_delete=models.CASCADE,
        related_name='documents',
        null=True,
        blank=True,
        verbose_name='Справа',
        help_text='Можна залишити порожнім для загальних документів'
    )
    
    category = models.ForeignKey(
        DocumentCategory,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='documents',
        verbose_name='Категорія'
    )
    
    document_type = models.CharField(
        max_length=30,
        choices=DOCUMENT_TYPE_CHOICES,
        verbose_name='Тип документу'
    )
    
    title = models.CharField(
        max_length=255,
        verbose_name='Назва документу'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Опис'
    )
    
    file = models.FileField(
        upload_to='documents/%Y/%m/',
        verbose_name='Файл'
    )
    
    file_size = models.IntegerField(
        editable=False,
        null=True,
        verbose_name='Розмір файлу (байт)'
    )
    
    document_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Дата документу',
        help_text='Дата створення/отримання документу'
    )
    
    # Метаданные
    author_name = models.CharField(
        max_length=255,
        blank=True,
        verbose_name='Автор документу',
        help_text='ПІБ автора документу (якщо відомо)'
    )
    
    recipient_name = models.CharField(
        max_length=255,
        blank=True,
        verbose_name='Одержувач',
        help_text='Кому адресовано документ'
    )
    
    registration_number = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Реєстраційний номер',
        help_text='Вхідний/вихідний номер документу'
    )
    
    # Full-Text Search
    search_vector = SearchVectorField(
        null=True,
        blank=True
    )
    
    # Права доступа
    is_confidential = models.BooleanField(
        default=True,
        verbose_name='Конфіденційний',
        help_text='Чи містить документ конфіденційну інформацію'
    )
    
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='uploaded_documents',
        verbose_name='Завантажив'
    )
    
    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Дата завантаження')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Дата оновлення')
    
    class Meta:
        verbose_name = 'Документ'
        verbose_name_plural = 'Документи'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['case', 'document_type']),
            models.Index(fields=['document_date']),
            GinIndex(fields=['search_vector']),  # PostgreSQL GIN index для Full-Text Search
        ]
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        # Сохраняем размер файла
        if self.file:
            self.file_size = self.file.size
        
        super().save(*args, **kwargs)
        
        # Обновляем search_vector для полнотекстового поиска
        if self.pk:
            Document.objects.filter(pk=self.pk).update(
                search_vector=(
                    SearchVector('title', weight='A') +
                    SearchVector('description', weight='B') +
                    SearchVector('registration_number', weight='C')
                )
            )
    
    def get_file_extension(self):
        """Получить расширение файла"""
        import os
        return os.path.splitext(self.file.name)[1].lower()
    
    @property
    def file_size_mb(self):
        """Размер файла в МБ"""
        if self.file_size:
            return round(self.file_size / (1024 * 1024), 2)
        return 0


class DocumentTemplate(models.Model):
    """
    Шаблоны документов для автоматической генерации.
    Используется python-docx для .docx файлов.
    """
    
    name = models.CharField(
        max_length=255,
        verbose_name='Назва шаблону'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Опис'
    )
    
    document_type = models.CharField(
        max_length=30,
        choices=Document.DOCUMENT_TYPE_CHOICES,
        verbose_name='Тип документу'
    )
    
    template_file = models.FileField(
        upload_to='templates/',
        verbose_name='Файл шаблону',
        help_text='Файл .docx з плейсхолдерами {{variable_name}}'
    )
    
    variables_help = models.TextField(
        blank=True,
        verbose_name='Опис змінних',
        help_text='Список доступних змінних та їх опис'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Активний'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        verbose_name='Створив'
    )
    
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        verbose_name = 'Шаблон документу'
        verbose_name_plural = 'Шаблони документів'
        ordering = ['name']
    
    def __str__(self):
        return self.name
    
    def generate_document(self, context_data):
        """
        Генерация документа на основе шаблона.
        context_data: словарь с данными для подстановки.
        """
        from docx import Document as DocxDocument
        from io import BytesIO
        
        # Открываем шаблон
        doc = DocxDocument(self.template_file.path)
        
        # Замена плейсхолдеров в параграфах
        for paragraph in doc.paragraphs:
            for key, value in context_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in context_data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
        
        # Сохраняем в BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output


# Регистрация моделей для аудита
auditlog.register(Document)
auditlog.register(DocumentTemplate)

